#!/usr/bin/env python3
"""Generate the FundExecs OS — GitHub org-migration playbook as a .docx.

Self-contained team-review document (no external dependencies beyond
python-docx). Mirrors docs/SECURITY_SETUP.md §3 with extra framing so it can be
handed to whoever runs the transfer.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x0A, 0x0F, 0x1C)
GOLD = RGBColor(0xC7, 0x8A, 0x12)
SLATE = RGBColor(0x47, 0x55, 0x69)

doc = Document()

# Base style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def heading(text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY if level <= 1 else SLATE
    return h


def para(text="", bold=False, italic=False, size=11, color=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(text, checkbox=False):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(("☐  " if checkbox else "") + text)
    p.paragraph_format.space_after = Pt(3)
    return p


def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    p.paragraph_format.space_after = Pt(3)
    return p


def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x0B, 0x3B, 0x2E)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(6)
    return p


# ── Cover ─────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
tr = title.add_run("FundExecs OS")
tr.bold = True
tr.font.size = Pt(26)
tr.font.color.rgb = NAVY
sub = doc.add_paragraph()
sr = sub.add_run("GitHub Org-Migration Playbook")
sr.bold = True
sr.font.size = Pt(16)
sr.font.color.rgb = GOLD
para(
    'Moving the repository into a GitHub Enterprise org — the "clone details over" plan.',
    italic=True,
    color=SLATE,
)
para("Prepared 2026-06-07 · for team review · repo: bey-group-international/fundexecs-os", size=9, color=SLATE)

doc.add_paragraph()

# ── Read-me-first ──────────────────────────────────────────────────────────────
heading("Before you start", 1)
para(
    "This is an owner-run operation in GitHub's UI — it cannot be automated from the "
    "codebase or from a repo-scoped agent. Everything in the repo (CI, CODEOWNERS, "
    "Dependabot, branch ruleset, CodeQL) is already config-as-code, so a new org "
    "inherits identical guardrails once you re-point integrations. Budget ~30–60 "
    "minutes and a short merge freeze."
)
para(
    "Goal: move history, issues, and PRs intact; lose nothing; keep CI / previews / "
    "review automation working on the other side.",
    italic=True,
    color=SLATE,
)

# ── Step 1: choose method ──────────────────────────────────────────────────────
heading("Step 1 — Choose the method", 1)
heading("Option A — Transfer (recommended if you are MOVING)", 2)
bullet("Preserves git history, issues, pull requests, releases, wiki, stars, and watchers.")
bullet("Sets up automatic redirects from the old repo URL to the new one.")
bullet("Done from: Repo → Settings → Danger Zone → Transfer.")
heading("Option B — Mirror (if you are DUPLICATING, keeping the original)", 2)
para("Create the empty target repo first, then from a terminal:")
code("git clone --mirror https://github.com/bey-group-international/fundexecs-os.git")
code("cd fundexecs-os.git && git push --mirror https://github.com/<NEW-ORG>/fundexecs-os.git")
para(
    "Important: a mirror copies git data ONLY. Issues, PRs, settings, and "
    "integrations do NOT come along — you would recreate those by hand.",
    bold=True,
)

# ── Step 2: pre-transfer ───────────────────────────────────────────────────────
heading("Step 2 — Pre-transfer checklist", 1)
bullet("Confirm you are an Owner on BOTH the source and destination orgs.", checkbox=True)
bullet("Announce a short merge freeze so no PRs are in flight during the move.", checkbox=True)
bullet(
    "Inventory the integrations to re-wire (Step 4) — Vercel, Supabase, CodeRabbit, "
    "any other GitHub Apps.",
    checkbox=True,
)
bullet(
    "List your Actions secrets + variables by NAME (values do not transfer): "
    "SUPABASE_*, STRIPE_*, ANTHROPIC_*, VOYAGE_*, ENABLE_CODEQL, etc.",
    checkbox=True,
)
bullet("Make sure the destination org's plan supports private-repo rulesets (Team/Enterprise).", checkbox=True)

# ── Step 3: transfer ───────────────────────────────────────────────────────────
heading("Step 3 — Run the transfer", 1)
numbered("Repo → Settings → Danger Zone → Transfer ownership.")
numbered("Enter the destination org and confirm the repository name.")
numbered("GitHub moves history / issues / PRs and installs URL redirects from the old path.")
numbered("Verify the repo opens under the new org and the default branch is intact.")

# ── Step 4: post-transfer rewire ───────────────────────────────────────────────
heading("Step 4 — Post-transfer re-wire (the part people forget)", 1)
para("Work top to bottom; each is quick but skipping one breaks CI or deploys.", italic=True, color=SLATE)
bullet(
    "Actions secrets + variables — recreate every secret/variable in the new repo "
    "(values are NOT carried over). Re-add ENABLE_CODEQL once Advanced Security is on.",
    checkbox=True,
)
bullet(
    "Default workflow token — set to read-only (Settings → Actions → General → "
    "Workflow permissions); our workflows request write scopes explicitly.",
    checkbox=True,
)
bullet(
    "Vercel — reconnect the Git repo (Vercel → Project → Settings → Git) so preview "
    "and production deploys keep firing; re-add env vars if the project is recreated.",
    checkbox=True,
)
bullet(
    "Supabase — reconnect the GitHub branch integration (Supabase → Integrations → "
    "GitHub) for preview branches and migration runs.",
    checkbox=True,
)
bullet(
    "CodeRabbit — install the CodeRabbit GitHub App on the new org "
    "(.coderabbit.yaml travels with the repo).",
    checkbox=True,
)
bullet("Other GitHub Apps / OAuth — reinstall on the new org as needed.", checkbox=True)
bullet(
    "Branch ruleset — re-import .github/rulesets/main-branch-protection.json "
    "(rulesets are per-org and are NOT carried by transfer).",
    checkbox=True,
)
bullet(
    "CODEOWNERS — already lists the BGI owners by email (pres@, "
    "businessdevelopment@, vp@, secretary@); confirm each maps to a GitHub "
    "account with repo access in the new org, or switch to team handles.",
    checkbox=True,
)
bullet(
    "Advanced Security — re-enable Secret Protection + Code Security on the repo in "
    "the new org, then set ENABLE_CODEQL=true to wake the CodeQL workflow.",
    checkbox=True,
)
bullet(
    "Hardcoded URLs — update README badges, docs, and any automation remotes that "
    "point at the old org path.",
    checkbox=True,
)

# ── Step 5: verify ─────────────────────────────────────────────────────────────
heading("Step 5 — Verify", 1)
bullet("Open a throwaway PR — confirm CI ('Typecheck, lint & build' + 'Playwright e2e smoke') runs and the ruleset blocks merge until green + approved.", checkbox=True)
bullet("Confirm the Vercel preview deploy posts on the PR.", checkbox=True)
bullet("Confirm CodeRabbit reviews the PR.", checkbox=True)
bullet("Confirm Dependabot is enabled (Settings → Code security).", checkbox=True)
bullet("Close the throwaway PR.", checkbox=True)

# ── What carries over (reference table) ────────────────────────────────────────
heading("Reference — what carries over", 1)
table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Item"
hdr[1].text = "Transfer (Option A)"
hdr[2].text = "Mirror (Option B)"
rows = [
    ("Git history / branches / tags", "Yes", "Yes"),
    ("Issues & pull requests", "Yes", "No"),
    ("Releases / wiki / stars", "Yes", "No"),
    ("Old-URL redirects", "Yes", "No"),
    ("Actions secrets & variables", "No (recreate)", "No (recreate)"),
    ("Branch rulesets / settings", "No (re-import)", "No (re-import)"),
    ("Integrations (Vercel/Supabase/CodeRabbit)", "No (reconnect)", "No (reconnect)"),
]
for item, a, b in rows:
    cells = table.add_row().cells
    cells[0].text = item
    cells[1].text = a
    cells[2].text = b
for r in table.rows:
    for c in r.cells:
        for p in c.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9.5)

doc.add_paragraph()
para(
    "Note: this playbook is a companion to docs/SECURITY_SETUP.md (repo hardening + "
    "Advanced Security enablement). Migrate only after the security add-ons and "
    "ruleset are decided, so the new org starts hardened.",
    italic=True,
    color=SLATE,
    size=10,
)

out = "docs/FundExecs_GitHub_Migration_Playbook.docx"
doc.save(out)
print("wrote", out)
